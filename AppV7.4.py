import os
import json
import io
import time
import getpass
import logging
import subprocess
import platform

import pickle
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_BREAK
import json
from bs4 import BeautifulSoup
import markdown
import shutil
from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException,StaleElementReferenceException
from colorama import init, Fore, Style
import zipfile
from rich.console import Console
from rich.table import Table
from rich.progress import Progress, SpinnerColumn, BarColumn, TextColumn
from PIL import Image
from AiProcessing import AiProcessing



init(autoreset=True)

CELIGO_ASCII = """

                                
        █████╗ ██╗   ██╗████████╗ ██████╗ ███╗   ██╗ ██████╗ ███╗   ███╗ ██████╗ ███████╗
        ██╔══██╗██║   ██║╚══██╔══╝██╔═══██╗████╗  ██║██╔═══██╗████╗ ████║██╔═══██╗██╔════╝
        ███████║██║   ██║   ██║   ██║   ██║██╔██╗ ██║██║   ██║██╔████╔██║██║   ██║███████╗
        ██╔══██║██║   ██║   ██║   ██║   ██║██║╚██╗██║██║   ██║██║╚██╔╝██║██║   ██║╚════██║
        ██║  ██║╚██████╔╝   ██║   ╚██████╔╝██║ ╚████║╚██████╔╝██║ ╚═╝ ██║╚██████╔╝███████║
        ╚═╝  ╚═╝ ╚═════╝    ╚═╝    ╚═════╝ ╚═╝  ╚═══╝ ╚═════╝ ╚═╝     ╚═╝ ╚═════╝ ╚══════╝
                                                                                                
                Dev : @IDrissTalainte
                Version : V7.4.0 Beta                   

                                    
            "There’s no use talking about the problem unless you talk about the solution."
                                                                            -Betty Williams
                                            

                                        
"""


def resize_terminal(rows: int = 39, columns: int = 157):
    """
    Resize the terminal window for both macOS and Windows.
    
    :param rows: Number of rows for the terminal window (default 39)
    :param columns: Number of columns for the terminal window (default 157)
    """
    current_os = platform.system()

    try:
        if current_os == "Darwin":  # macOS
            applescript = f'''
            tell application "Terminal"
                set current settings of selected tab of front window to settings set "Basic"
                set number of rows of selected tab of front window to {rows}
                set number of columns of selected tab of front window to {columns}
            end tell
            '''
            subprocess.run(["osascript", "-e", applescript], check=True, capture_output=True, text=True)
        
        elif current_os == "Windows":
            os.system(f'mode con: cols={columns} lines={rows}')
        
        else:
            logging.warning(f"Terminal resizing not supported on {current_os}")
            return

        logging.info(f"Terminal resized to {rows} rows and {columns} columns on {current_os}")
        time.sleep(1)  # Wait for 1 second to ensure the resize has taken effect

    except subprocess.CalledProcessError as e:
        logging.error(f"Failed to resize terminal on {current_os}: {e.stderr}")
    except Exception as e:
        logging.error(f"Unexpected error while resizing terminal on {current_os}: {str(e)}")


class CeligoAutomation:
    def __init__(self):
        self.console = Console()
        self.driver = None
        self.project_dir = os.path.dirname(os.path.abspath(__file__))
        self.celigo_ai_dir = os.path.join(self.project_dir, "CeligoAI")
        self.output_directory = os.path.join(self.celigo_ai_dir, "DocumentResources")
        self.cookie_file = os.path.join(self.celigo_ai_dir, "celigo_cookies.pkl")
        self.display_ascii_art()
     


    def display_menu(self):
        menu = """
🤖 Celigo Automation Menu:
1. 🚀 Automate All (Recommended)
2. 📦 Extract Your Integration
3. 📸 Generate Screenshots
4. 🧠 Generate AI Descriptions
5. 📄 Generate O&M
6. 🧹 Clean All Generated Files
7. 🚪 Exit
        """
        self.console.print(menu)
    def capture_flow_screenshot(self, integration_id, flow_id):
        flow_url = f"https://integrator.io/integrations/{integration_id}/flowBuilder/{flow_id}"
        self.driver.get(flow_url)

        try:
            # Wait for the react-flow__pane element to be available
            pane_element = WebDriverWait(self.driver, 20).until(
                EC.presence_of_element_located((By.CLASS_NAME, "react-flow__pane"))
            )

            # Wait for 7 seconds to allow the flow to fully render
            time.sleep(13)

            # Find and click the zoom fit button using aria-label
            zoom_fit_button = self.driver.find_element(By.XPATH, "//button[./span[@aria-label='Zoom to fit']]")
            ActionChains(self.driver).move_to_element(zoom_fit_button).click().perform()

            # Wait a bit for the zoom animation to complete
            time.sleep(3)

            # Capture the screenshot of the react-flow__pane element
            screenshot = pane_element.screenshot_as_png
            
            # Save the screenshot
            filename = f"{flow_id}_flow.png"
            filepath = os.path.join(self.output_directory, filename)
            with open(filepath, "wb") as file:
                file.write(screenshot)

            self.log(f"Captured flow screenshot: {filename}", color=Fore.GREEN)
            return filename
        except Exception as e:
            self.log(f"Error capturing flow screenshot for flow {flow_id}: {str(e)}", level="ERROR", color=Fore.RED)
            return None
        
    def process_scripts(self):
        self.log("Starting script processing...", color=Fore.CYAN)
        extracted_dir = os.path.join(self.celigo_ai_dir, "extracted")
        scripts_dir = os.path.join(extracted_dir, "scripts")
        
        if not os.path.exists(scripts_dir):
            self.log(f"Scripts directory not found: {scripts_dir}", level="ERROR", color=Fore.RED)
            return

        if not os.path.exists(self.output_directory):
            os.makedirs(self.output_directory)
            self.log(f"Created output directory: {self.output_directory}", color=Fore.GREEN)

        def get_short_id(full_id):
            return full_id[-24:]

        def find_script_ids(obj):
            script_ids = []
            if isinstance(obj, dict):
                for key, value in obj.items():
                    if key == '_scriptId':
                        script_ids.append(value)
                    elif isinstance(value, (dict, list)):
                        script_ids.extend(find_script_ids(value))
            elif isinstance(obj, list):
                for item in obj:
                    script_ids.extend(find_script_ids(item))
            return script_ids

        scripts_processed = 0
        for folder in ["exports", "imports"]:
            folder_path = os.path.join(extracted_dir, folder)
            if os.path.exists(folder_path):
                files = [f for f in os.listdir(folder_path) if f.endswith('.json')]
                self.log(f"Processing {len(files)} files in {folder}", color=Fore.CYAN)
                for file in files:
                    file_path = os.path.join(folder_path, file)
                    try:
                        with open(file_path, 'r') as f:
                            data = json.load(f)
                        
                        script_ids = find_script_ids(data)
                        
                        if script_ids:
                            for script_id in script_ids:
                                # Find the corresponding script file
                                script_files = [f for f in os.listdir(scripts_dir) if f.endswith(get_short_id(script_id) + '.js')]
                                if script_files:
                                    script_file = script_files[0]
                                    script_path = os.path.join(scripts_dir, script_file)
                                    
                                    # Get the import/export ID
                                    item_id = get_short_id(os.path.splitext(file)[0])
                                    
                                    # Create the new file name
                                    new_script_name = f"{item_id}_script_{scripts_processed + 1}.txt"
                                    new_script_path = os.path.join(self.output_directory, new_script_name)
                                    
                                    # Copy the file
                                    shutil.copy2(script_path, new_script_path)
                                    self.log(f"Copied script for {folder[:-1]} {item_id} to {new_script_path}", color=Fore.GREEN)
                                    scripts_processed += 1
                                else:
                                    self.log(f"Script file not found for ID: {script_id}", level="WARNING", color=Fore.YELLOW)
                        else:
                            self.log(f"No scriptId found for {folder[:-1]} {get_short_id(os.path.splitext(file)[0])}", level="INFO", color=Fore.YELLOW)
                    except Exception as e:
                        self.log(f"Error processing {file}: {str(e)}", level="ERROR", color=Fore.RED)

        self.log(f"Script processing completed. Processed {scripts_processed} scripts.", color=Fore.GREEN)
        self.log(f"Scripts saved in: {self.output_directory}", color=Fore.GREEN)
    def add_image_to_document(doc, image_path, caption):
        # Add the image
        doc.add_picture(image_path, width=Inches(6))
        
        # Add a paragraph for the image and center it
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add the caption as a separate paragraph
        caption_paragraph = doc.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.style = doc.styles['Caption']
    def generate_documentation(self):
        self.log("Starting documentation generation...", color=Fore.CYAN)
        doc_name = f"O&M"
        
        doc = Document()
        
        # Define styles
        styles = doc.styles
        style_heading1 = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        style_heading1.font.name = 'Calibri'
        style_heading1.font.size = Pt(18)
        style_heading1.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        style_heading1.font.bold = True
        
        style_heading2 = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        style_heading2.font.name = 'Calibri'
        style_heading2.font.size = Pt(16)
        style_heading2.font.color.rgb = RGBColor(0, 102, 204)  # Medium blue
        style_heading2.font.bold = True
        
        style_heading3 = styles.add_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH)
        style_heading3.font.name = 'Calibri'
        style_heading3.font.size = Pt(14)
        style_heading3.font.color.rgb = RGBColor(51, 153, 255)  # Light blue
        style_heading3.font.bold = True
        
        style_normal = styles.add_style('Custom Normal', WD_STYLE_TYPE.PARAGRAPH)
        style_normal.font.name = 'Calibri'
        style_normal.font.size = Pt(11)
        style_normal.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        style_code = styles.add_style('Code Style', WD_STYLE_TYPE.PARAGRAPH)
        style_code.font.name = 'Consolas'
        style_code.font.size = Pt(10)
        style_code.font.color.rgb = RGBColor(0, 102, 0)  # Dark green
        
        # Add title
        title = doc.add_heading('LYSI', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def find_matching_files(directory):
            all_files = os.listdir(directory)
            flows = sorted([f for f in all_files if f.endswith('_flow.png')])
            images = sorted([f for f in all_files if (f.startswith(('export_', 'import_')) and f.lower().endswith(('.png', '.jpg', '.jpeg')))])
            descriptions = [f for f in all_files if f.endswith('_Description.md')]
            scripts = [f for f in all_files if f.endswith('_script_1.txt')]

            self.log(f"Found {len(flows)} flows, {len(images)} images, {len(descriptions)} descriptions, and {len(scripts)} scripts.", color=Fore.CYAN)

            matched_files = []

            # Add flows first
            for flow in flows:
                flow_id = flow.split('_flow.png')[0][-24:]
                matched_files.append({
                    'type': 'flow',
                    'image': flow,
                    'id': flow_id,
                })

            # Then add exports and imports
            for image in images:
                full_id = image.split('_', 1)[1].split('.')[0]
                short_id_10 = full_id[-10:]
                short_id_24 = full_id[-24:]
                
                matching_description = next((d for d in descriptions if d.split('_Description.md')[0][-10:] == short_id_10), None)
                matching_script = next((s for s in scripts if s.split('_script_1.txt')[0][-10:] == short_id_10), None)

                matched_files.append({
                    'type': 'export' if image.startswith('export_') else 'import',
                    'image': image,
                    'description': matching_description,
                    'script': matching_script,
                    'id_10': short_id_10,
                    'id_24': short_id_24,
                })

            return matched_files

        def get_title_from_json(item_type, item_id):
            folder_name = f"{item_type}s"  # 'flows', 'exports', or 'imports'
            json_folder = os.path.join(self.celigo_ai_dir, "extracted", folder_name)
            json_files = [f for f in os.listdir(json_folder) if f.endswith('.json')]
            
            for json_file in json_files:
                if json_file.split('.json')[0][-24:] == item_id:
                    json_path = os.path.join(json_folder, json_file)
                    with open(json_path, 'r') as f:
                        data = json.load(f)
                        return data.get('name', 'Unnamed Item')
            return 'Unnamed Item'

        def convert_markdown_to_docx(markdown_text, doc):
            html = markdown.markdown(markdown_text, extensions=['tables'])
            soup = BeautifulSoup(html, 'html.parser')
            
            for element in soup.find_all():
                if element.name in ['h1', 'h2', 'h3']:
                    level = int(element.name[1])
                    p = doc.add_paragraph(element.text, style=f'Custom Heading {level}')
                elif element.name == 'p':
                    p = doc.add_paragraph(style='Custom Normal')
                    for child in element.children:
                        if child.name == 'strong':
                            p.add_run(child.text).bold = True
                        elif child.name == 'em':
                            p.add_run(child.text).italic = True
                        elif child.name == 'u':
                            p.add_run(child.text).underline = True
                        else:
                            p.add_run(child.text)
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li'):
                        style = 'List Bullet' if element.name == 'ul' else 'List Number'
                        p = doc.add_paragraph(style=style)
                        for child in li.children:
                            if child.name == 'strong':
                                p.add_run(child.text).bold = True
                            elif child.name == 'em':
                                p.add_run(child.text).italic = True
                            elif child.name == 'u':
                                p.add_run(child.text).underline = True
                            else:
                                p.add_run(child.text)
                elif element.name == 'pre':
                    code = element.find('code')
                    if code:
                        p = doc.add_paragraph(style='Code Style')
                        p.add_run(code.text)
                elif element.name == 'table':
                    rows = element.find_all('tr')
                    table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['th', 'td'])))
                    table.style = 'Table Grid'
                    
                    for i, row in enumerate(rows):
                        cells = row.find_all(['th', 'td'])
                        for j, cell in enumerate(cells):
                            table_cell = table.cell(i, j)
                            for child in cell.children:
                                if child.name == 'strong':
                                    table_cell.paragraphs[0].add_run(child.text).bold = True
                                elif child.name == 'em':
                                    table_cell.paragraphs[0].add_run(child.text).italic = True
                                elif child.name == 'u':
                                    table_cell.paragraphs[0].add_run(child.text).underline = True
                                else:
                                    table_cell.paragraphs[0].add_run(child.text)
                            
                            if cell.name == 'th':
                                for paragraph in table_cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                    
                    doc.add_paragraph()  # Add space after the table

        matched_files = find_matching_files(self.output_directory)

        # Process flows first
        flow_files = [f for f in matched_files if f['type'] == 'flow']
        for flow in flow_files:
            try:
                image_file = flow['image']
                item_id = flow['id']

                # Get title from JSON
                title = get_title_from_json('flow', item_id)

                # Add title
                doc.add_heading(f"Flow: {title}", level=1)

                # Add image
                image_path = os.path.join(self.output_directory, image_file)
                doc.add_picture(image_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.log(f"Added flow image: {image_file}", color=Fore.GREEN)

                # Add a page break after each flow
                doc.add_page_break()

            except Exception as e:
                self.log(f"Error processing flow {flow['image']}: {str(e)}", level="ERROR", color=Fore.RED)

        # Process exports and imports
        for match in [f for f in matched_files if f['type'] != 'flow']:
            try:
                item_type = match['type']
                image_file = match['image']
                item_id_10 = match['id_10']
                item_id_24 = match['id_24']

                # Get title from JSON using 24-character ID
                title = get_title_from_json(item_type, item_id_24)

                # Add title
                doc.add_heading(f"{item_type.capitalize()}: {title}", level=1)

                # Add image
                image_path = os.path.join(self.output_directory, image_file)
                doc.add_picture(image_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.log(f"Added {item_type} image: {image_file}", color=Fore.GREEN)

                # Add description if exists (using 10-character ID)
                if 'description' in match and match['description']:
                    description_file = match['description']
                    description_path = os.path.join(self.output_directory, description_file)
                    with open(description_path, 'r') as desc_file:
                        description = desc_file.read()
                    convert_markdown_to_docx(description, doc)
                    self.log(f"Added formatted description for {image_file}", color=Fore.GREEN)
                else:
                    doc.add_paragraph("No description available for this item.", style='Custom Normal')
                    self.log(f"No description file found for {image_file}", level="WARNING", color=Fore.YELLOW)

                # Add script if exists (using 10-character ID)
                if 'script' in match and match['script']:
                    script_file = match['script']
                    script_path = os.path.join(self.output_directory, script_file)
                    doc.add_heading("Script", level=2)
                    with open(script_path, 'r') as script_file:
                        script_content = script_file.read()
                    p = doc.add_paragraph(style='Code Style')
                    p.add_run(script_content)
                    self.log(f"Added script for {image_file}", color=Fore.GREEN)

                # Add a page break after each export/import
                doc.add_page_break()

            except Exception as e:
                self.log(f"Error processing {match['image']}: {str(e)}", level="ERROR", color=Fore.RED)

        # Add footer
        section = doc.sections[-1]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "LYSI Consulting"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save the document
        doc_path = os.path.join(self.output_directory, f"{doc_name}.docx")
        doc.save(doc_path)
        self.log(f"Documentation generated and saved to: {doc_path}", color=Fore.GREEN)

    def log(self, message, level="INFO", color=Fore.WHITE):
        timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
        formatted_message = f"[{timestamp}] [{level}] {message}"
        print(f"{color}{formatted_message}{Style.RESET_ALL}")
    
    def display_ascii_art(self):
        resize_terminal()
        self.console.print(CELIGO_ASCII)
   
    def get_webdriver(self):
        self.log("Hi Develper I know you feel Overwhelmed i'm here to help you (: ) ", color=Fore.CYAN)
        options = ChromeOptions()
        options.add_argument("--start-maximized")
        options.add_argument("---headless")
        options.add_argument("--disable-infobars")
        options.add_argument("--disable-extensions")
        options.add_argument("--disable-popup-blocking")
        service = ChromeService()
        self.driver = webdriver.Chrome(service=service, options=options)
        self.log("Let's Goo!.", color=Fore.GREEN)

    def save_cookies(self):
        if self.driver:
            self.log("Saving cookies...", color=Fore.CYAN)
            cookies = self.driver.get_cookies()
            with open(self.cookie_file, 'wb') as f:
                pickle.dump(cookies, f)
            self.log("Cookies saved successfully.", color=Fore.GREEN)

    def load_cookies(self):
        if os.path.exists(self.cookie_file):
            self.log("Loading saved cookies...", color=Fore.CYAN)
            with open(self.cookie_file, 'rb') as f:
                cookies = pickle.load(f)
            for cookie in cookies:
                if 'expiry' in cookie:
                    del cookie['expiry']
                self.driver.add_cookie(cookie)
            self.log("Cookies loaded successfully.", color=Fore.GREEN)
            return True
        else:
            self.log("No saved cookies found.", color=Fore.YELLOW)
            return False

    def perform_login(self):
        self.log("Starting login process...", color=Fore.CYAN)
        self.driver.get("https://integrator.io")
        self.load_cookies()
        time.sleep(5)  # Wait for cookies to take effect
        if "/signin" in self.driver.current_url:
            return self.manual_login()
        elif self.check_login_status():
            self.log("Successfully logged in using cookies.", color=Fore.GREEN)
            self.close_pendo_guide()  # New method to close the Pendo guide
            return True
        else:
            self.log("Unexpected state after loading cookies.", color=Fore.YELLOW)
            return self.manual_login()

    def close_pendo_guide(self):
        try:
            close_button = WebDriverWait(self.driver, 10).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "[id^='pendo-close-guide-']"))
            )
            close_button.click()
            self.log("Closed Pendo guide successfully.", color=Fore.GREEN)
        except TimeoutException:
            self.log("", color=Fore.YELLOW)
        except Exception as e:
            self.log(f"", color=Fore.RED)

    def manual_login(self):
        email = input("Enter your Celigo email: ")
        password = getpass.getpass("Enter your Celigo password: ")
        self.driver.get("https://integrator.io/signin")
        try:
            email_input = self.wait_for_element(By.CSS_SELECTOR, "input[type='email']")
            email_input.clear()
            email_input.send_keys(email)
            password_input = self.wait_for_element(By.CSS_SELECTOR, "input[type='password']")
            password_input.clear()
            password_input.send_keys(password)
            sign_in_button = self.wait_for_element(By.CSS_SELECTOR, "button[type='submit']")
            sign_in_button.click()
            time.sleep(7)  # Wait for response
            if "/mfa/verify" in self.driver.current_url:
                return self.handle_2fa()
            elif self.check_login_status():
                self.log("Successfully logged in to Celigo.", color=Fore.GREEN)
                self.save_cookies()
                return True
            else:
                self.log("Login failed. Please check your credentials.", color=Fore.RED)
                return False
        except Exception as e:
            self.log(f"Error during manual login: {str(e)}", color=Fore.RED)
            return False

    def handle_2fa(self):
        max_attempts = 3
        for attempt in range(max_attempts):
            verification_code = input("Enter your 2FA code: ")
            try:
                code_input = self.wait_for_element(By.CSS_SELECTOR, "input[type='text']")
                code_input.clear()
                code_input.send_keys(verification_code)
                submit_button = self.wait_for_element(By.CSS_SELECTOR, "button[type='submit']")
                submit_button.click()
                self.log(f"Submitted 2FA code (Attempt {attempt + 1}). Waiting for navigation...", color=Fore.CYAN)
                time.sleep(5)  # Wait for navigation
                if self.check_login_status():
                    self.log("Successfully logged in to Celigo after 2FA.", color=Fore.GREEN)
                    self.save_cookies()
                    return True
            except Exception as e:
                self.log(f"Error during 2FA: {str(e)}", color=Fore.RED)
            self.log("2FA verification failed. Please try again.", color=Fore.YELLOW)
        self.log("Failed to verify 2FA after multiple attempts.", color=Fore.RED)
        return False

    def check_login_status(self):
        try:
            WebDriverWait(self.driver, 10).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, "h3.MuiTypography-root.MuiTypography-h3"))
            )
            self.log("Login confirmed. User is on the home page.", color=Fore.GREEN)
            return True
        except TimeoutException:
            self.log(f"Login status check failed. Current URL: {self.driver.current_url}", color=Fore.YELLOW)
            return False

    def wait_for_element(self, by, value, timeout=10):
        try:
            element = WebDriverWait(self.driver, timeout).until(
                EC.presence_of_element_located((by, value))
            )
            return element
        except TimeoutException:
            self.log(f"Timeout waiting for element: {by}={value}", level="ERROR", color=Fore.RED)
            return None


    def resize_terminal(rows: int = 39, columns: int = 157):
        try:
            subprocess.run(["osascript", "-e", applescript], check=True, capture_output=True, text=True)
            logging.info(f"Terminal resized to {rows} rows and {columns} columns.")
        except subprocess.CalledProcessError as e:
            logging.error(f"Failed to resize terminal: {e.stderr}")
    def select_project(self):
            self.close_pendo_guide()  # Ensure Pendo guide is closed before selecting project
            max_retries = 3

            for attempt in range(max_retries):
                try:
                    self.log("Selecting project...", color=Fore.CYAN)
                    project_button = self.wait_for_element(
                        By.CSS_SELECTOR, "button.MuiButtonBase-root.MuiButton-root.MuiButton-text.MuiButton-textSecondary"
                    )
                    project_button.click()
                    project_list = self.wait_for_element(
                        By.CSS_SELECTOR, "ul.MuiList-root.MuiList-padding.MuiList-dense"
                    )
                    project_items = project_list.find_elements(By.CSS_SELECTOR, "li.MuiListItem-container")

                    table = Table(title="Available Projects")
                    table.add_column("Number", style="cyan", no_wrap=True)
                    table.add_column("Project Name", style="magenta")

                    projects = []
                    for i, item in enumerate(project_items, 1):
                        project_name = item.find_element(By.CSS_SELECTOR, "span.MuiTypography-root.MuiTypography-body2.MuiListItemText-primary").text
                        projects.append(project_name)
                        table.add_row(str(i), project_name)

                    self.console.print(table)

                    while True:
                        try:
                            choice = int(input("Enter the number of the project you want to select: "))
                            if 1 <= choice <= len(projects):
                                selected_project = projects[choice - 1]
                                # Find the project element again to avoid stale element reference
                                for item in self.driver.find_elements(By.CSS_SELECTOR, "li.MuiListItem-container"):
                                    project_span = item.find_element(By.CSS_SELECTOR, "span.MuiTypography-root.MuiTypography-body2.MuiListItemText-primary")
                                    if project_span.text == selected_project:
                                        project_span.click()
                                        self.log(f"Selected project: {selected_project}", color=Fore.GREEN)
                                        time.sleep(5)  # Wait for project to load
                                        return True
                                raise Exception(f"Project '{selected_project}' not found in the list.")
                            else:
                                print("Invalid choice. Please enter a number from the list.")
                        except ValueError:
                            print("Invalid input. Please enter a number.")

                except StaleElementReferenceException:
                    if attempt < max_retries - 1:
                        self.log("Stale element encountered. Retrying...", color=Fore.YELLOW)
                        time.sleep(2)  # Wait before retrying
                    else:
                        self.log("Failed to select project after multiple attempts.", level="ERROR", color=Fore.RED)
                        return False
                except Exception as e:
                    self.log(f"Error selecting project: {str(e)}", level="ERROR", color=Fore.RED)
                    return False

            return False
                 

    def select_environment(self):
        self.log("Selecting environment...", color=Fore.CYAN)
        env_table = Table(title="Select Environment")
        env_table.add_column("Option", style="cyan", no_wrap=True)
        env_table.add_column("Environment", style="magenta")
        env_table.add_row("1", "Production")
        env_table.add_row("2", "Sandbox")
        env_table.add_row("3", "No Environment")
        self.console.print(env_table)

        while True:
            env_choice = input("Enter the number of the environment you want to select (1, 2, or 3): ").strip()
            if env_choice == "1":
                environment = "Production"
                break
            elif env_choice == "2":
                environment = "Sandbox"
                break
            elif env_choice == "3":
                environment = "No Environment"
                break
            else:
                self.log("Invalid choice. Please enter 1 for Production, 2 for Sandbox, or 3 for No Environment.", color=Fore.RED)

        if environment == "No Environment":
            self.log("Proceeding without selecting an environment.", color=Fore.GREEN)
            return True

        button_selector = (
            "button.MuiButtonBase-root.MuiToggleButton-root.MuiToggleButton-sizeMedium"
            ".MuiToggleButton-standard.MuiToggleButtonGroup-grouped.MuiToggleButtonGroup-groupedHorizontal"
        )
        try:
            buttons = self.driver.find_elements(By.CSS_SELECTOR, button_selector)
            for button in buttons:
                if button.text.strip().lower() == environment.lower():
                    button.click()
                    self.log(f"Clicked {environment} environment button", color=Fore.GREEN)
                    time.sleep(3)  # Wait for the environment to switch
                    return True
            self.log(f"Could not find {environment} environment button", level="ERROR", color=Fore.RED)
            return False
        except Exception as e:
            self.log(f"Error selecting environment: {str(e)}", level="ERROR", color=Fore.RED)
            return False

    def process_integration_zip(self):
        self.log("Starting integration ZIP processing...", color=Fore.CYAN)
        integration_data = {"flows": [], "descriptions": {}}
        flow_ids = []  # Initialize flow_ids list

        def get_short_id(full_id):
            return full_id[-24:] if len(full_id) > 23 else full_id

        zip_files = [f for f in os.listdir(self.celigo_ai_dir) if f.endswith(".zip")]
        if not zip_files:
            self.log("No ZIP file found in the CeligoAI directory.", level="ERROR", color=Fore.RED)
            return None, None, []

        zip_file = zip_files[0]
        zip_path = os.path.join(self.celigo_ai_dir, zip_file)
        self.log(f"Found ZIP file: {zip_file}", color=Fore.GREEN)

        extract_dir = os.path.join(self.celigo_ai_dir, "extracted")
        self.log(f"Extracting ZIP file to: {extract_dir}", color=Fore.CYAN)
        with zipfile.ZipFile(zip_path, "r") as zip_ref:
            zip_ref.extractall(extract_dir)

        integration_id = os.path.splitext(zip_file)[0]
        self.log(f"Integration ID: {integration_id}", color=Fore.GREEN)

        os.makedirs(self.output_directory, exist_ok=True)

        flows_dir = os.path.join(extract_dir, "flows")
        self.log("Processing flows...", color=Fore.CYAN)
        flow_files = [f for f in os.listdir(flows_dir) if f.endswith(".json")]

        with Progress() as progress:
            task = progress.add_task("[cyan]Processing flows...", total=len(flow_files))
            for flow_file in flow_files:
                with open(os.path.join(flows_dir, flow_file), "r") as f:
                    flow_data = json.load(f)
                    full_flow_id = os.path.splitext(flow_file)[0]
                    flow_id = get_short_id(full_flow_id)
                    flow_ids.append(flow_id)  # Add flow_id to the list
                    integration_data["flows"].append({
                        "id": flow_id,
                        "grouping_id": get_short_id(flow_data.get("_flowGroupingId", "")) if flow_data.get("_flowGroupingId") else None,
                        "exports": [],
                        "imports": [],
                    })
                progress.update(task, advance=1)

        for folder in ["exports", "imports"]:
            folder_path = os.path.join(extract_dir, folder)
            if os.path.exists(folder_path):
                self.log(f"Processing {folder}...", color=Fore.CYAN)
                files = [f for f in os.listdir(folder_path) if f.endswith(".json")]

                with Progress() as progress:
                    task = progress.add_task(f"[cyan]Processing {folder}...", total=len(files))
                    for file in files:
                        with open(os.path.join(folder_path, file), "r") as f:
                            data = json.load(f)
                            full_item_id = os.path.splitext(file)[0]
                            item_id = get_short_id(full_item_id)

                            for flow in integration_data["flows"]:
                                if item_id.startswith(flow["id"][:2]):
                                    flow[folder].append(item_id)
                                    break

                            integration_data["descriptions"][item_id] = data.get("aiDescription", {}).get("detailed", "No detailed description available")
                        progress.update(task, advance=1)

        self.log(f"Processed {len(integration_data['flows'])} flows.", color=Fore.GREEN)
        for flow in integration_data["flows"]:
            self.log(f"Flow {flow['id']}: {len(flow['exports'])} exports, {len(flow['imports'])} imports", color=Fore.GREEN)

        return integration_id, integration_data, flow_ids

    def generate_urls(self, integration_id, integration_data):
        self.log("Generating URLs...", color=Fore.CYAN)
        urls = []
        for flow in integration_data["flows"]:
            flow_id = flow["id"]
            if flow.get("grouping_id"):
                flow_url = f"https://integrator.io/integrations/{integration_id}/flows/sections/{flow['grouping_id']}/flowBuilder/{flow_id}"
            else:
                flow_url = f"https://integrator.io/integrations/{integration_id}/flowBuilder/{flow_id}"
            
            for export_id in flow["exports"]:
                url = f"{flow_url}/edit/exports/{export_id}"
                urls.append(("export", url, export_id))
            for import_id in flow["imports"]:
                url = f"{flow_url}/edit/imports/{import_id}"
                urls.append(("import", url, import_id))
        
        self.log(f"Generated {len(urls)} URLs for exports and imports.", color=Fore.GREEN)
        return urls

    def take_screenshots(self, urls):
        self.log("Starting screenshot capture process...", color=Fore.CYAN)
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(complete_style="blue"),
            TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
            console=self.console
        ) as progress:
            task = progress.add_task("[cyan]Capturing screenshots...", total=len(urls))
            for i, (item_type, url, item_id) in enumerate(urls, 1):
                max_retries = 3
                for attempt in range(max_retries):
                    try:
                        self.driver.get(url)
                        time.sleep(15)  # Wait for the page to load
                        drawer_selector = "div.MuiPaper-root.MuiPaper-elevation.MuiPaper-elevation16.MuiDrawer-paper.MuiDrawer-paperAnchorRight"
                        drawer_element = self.wait_for_element(By.CSS_SELECTOR, drawer_selector, timeout=20)
                        if drawer_element:
                            filename = f"{item_type}_{item_id}.png"
                            filepath = os.path.join(self.output_directory, filename)
                            self.take_full_element_screenshot(drawer_element, filepath)
                            break
                        else:
                            raise NoSuchElementException("Drawer element not found")
                    except (TimeoutException, NoSuchElementException, StaleElementReferenceException) as e:
                        if attempt < max_retries - 1:
                            self.log(f"Attempt {attempt + 1} failed. Retrying... Error: {str(e)}", level="WARNING", color=Fore.YELLOW)
                            time.sleep(4)  # Wait before retrying
                        else:
                            self.log(f"Failed to capture screenshot for {item_type} {item_id} after {max_retries} attempts.", level="ERROR", color=Fore.RED)
                            self.log(f"Error: {str(e)}", level="ERROR", color=Fore.RED)
                progress.update(task, advance=1)


    def take_full_element_screenshot(self, element, output_path):
        try:
            # Function to capture screenshot
            def capture_screenshot():
                png = self.driver.get_screenshot_as_png()
                im = Image.open(io.BytesIO(png))
                element_dimensions = self.driver.execute_script(
                    """
                    var rect = arguments[0].getBoundingClientRect();
                    return {
                        top: rect.top,
                        left: rect.left,
                        width: rect.width,
                        height: arguments[0].scrollHeight
                    };
                """,
                    element,
                )
                device_pixel_ratio = self.driver.execute_script(
                    "return window.devicePixelRatio;"
                )
                left = int(element_dimensions["left"] * device_pixel_ratio)
                top = int(element_dimensions["top"] * device_pixel_ratio)
                right = int(
                    (element_dimensions["left"] + element_dimensions["width"])
                    * device_pixel_ratio
                )
                bottom = int(
                    (element_dimensions["top"] + element_dimensions["height"])
                    * device_pixel_ratio
                )
                return im.crop((left, top, right, bottom))

            # Take first screenshot with all sections expanded
            expanded_screenshot = capture_screenshot()

            # Find the first expandable section
            expandable_sections = self.driver.find_elements(
                By.CSS_SELECTOR,
                "div.MuiButtonBase-root.MuiAccordionSummary-root.Mui-expanded.MuiAccordionSummary-gutters",
            )
            
            if expandable_sections:
                # Click to collapse only the first section
                self.driver.execute_script("arguments[0].click();", expandable_sections[0])
                time.sleep(1)  # Wait for collapse animation

                # Take second screenshot with first section collapsed
                collapsed_screenshot = capture_screenshot()

                # Merge the two screenshots
                total_height = expanded_screenshot.height + collapsed_screenshot.height
                merged_image = Image.new("RGB", (expanded_screenshot.width, total_height))
                merged_image.paste(expanded_screenshot, (0, 0))
                merged_image.paste(collapsed_screenshot, (0, expanded_screenshot.height))

                # Save the merged image
                merged_image.save(output_path)
            else:
                # If no expandable sections found, just save the expanded screenshot
                expanded_screenshot.save(output_path)

            self.log(f"Screenshot saved: {output_path}", color=Fore.GREEN)

        except Exception as e:
            self.log(
                f"Failed to take full element screenshot: {str(e)}",
                level="ERROR",
                color=Fore.RED,
            )
            raise

    def clean_integration_data(self):
        self.log("Starting data cleaning process...", color=Fore.CYAN)
        extracted_dir = os.path.join(self.celigo_ai_dir, "extracted")
        
        keys_to_remove = [
            "adaptorType", "distributed", "mockOutput","_id","paging", "_connectionId", "apiIdentifier",
            "asynchronous", "oneToMany", "type", "skipGrouping", "isRest",
            "resourcePath", "response", "restletVersion"
        ]

        for folder in ["exports", "imports"]:
            folder_path = os.path.join(extracted_dir, folder)
            if os.path.exists(folder_path):
                files = [f for f in os.listdir(folder_path) if f.endswith('.json')]
                with Progress() as progress:
                    task = progress.add_task(f"[cyan]Cleaning {folder}...", total=len(files))
                    for file in files:
                        file_path = os.path.join(folder_path, file)
                        with open(file_path, 'r') as f:
                            data = json.load(f)
                        cleaned_data = self.remove_keys(data, keys_to_remove)
                        with open(file_path, 'w') as f:
                            json.dump(cleaned_data, f, indent=2)
                        progress.update(task, advance=1)
        self.log("Data cleaning complete.", color=Fore.GREEN)

    def remove_keys(self, obj, keys):
        if isinstance(obj, dict):
            return {k: self.remove_keys(v, keys) for k, v in obj.items() if k not in keys}
        elif isinstance(obj, list):
            return [self.remove_keys(item, keys) for item in obj]
        else:
            return obj

    
    def generate_ai_descriptions(self):
        self.log("Starting AI description generation...", color=Fore.CYAN)
        try:
            ai_processor = AiProcessing(self.celigo_ai_dir)
            ai_processor.process_descriptions()
            self.log("AI description generation complete.", color=Fore.GREEN)
        except Exception as e:
            self.log(f"Error in AI description generation: {str(e)}", level="ERROR", color=Fore.RED)
   
    def perform_automation(self):
        #self.display_ascii_art()
        max_attempts = 5  
        
        def execute_with_retry(func, *args, **kwargs):
            for attempt in range(max_attempts):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    if attempt < max_attempts - 1:
                        self.log(f"Attempt {attempt + 1} failed. Retrying... Error: {str(e)}", level="WARNING", color=Fore.YELLOW)
                        time.sleep(5)  # Wait before retrying
                    else:
                        self.log(f"Failed after {max_attempts} attempts. Error: {str(e)}", level="ERROR", color=Fore.RED)
                        return False
            return False

        try:
            self.get_webdriver()
            login_success = False
            for attempt in range(max_attempts):
                login_success = execute_with_retry(self.perform_login)
                if login_success:
                    break
                elif "/signin" in self.driver.current_url:
                    self.log(f"Login attempt {attempt + 1} failed. Please try again.", level="WARNING", color=Fore.YELLOW)
                else:
                    break

            if not login_success:
                self.log("Login failed after multiple attempts. Exiting.", level="ERROR", color=Fore.RED)
                return

            project_selected = execute_with_retry(self.select_project)
            if not project_selected:
                self.log("Failed to select project. Exiting.", level="ERROR", color=Fore.RED)
                return

            environment_selected = execute_with_retry(self.select_environment)
            if environment_selected is None:
                self.log("Failed to select environment. Exiting.", level="ERROR", color=Fore.RED)
                return

            self.log("Login successful. Project and environment (if applicable) selected.", color=Fore.GREEN)

            while True:
                self.display_menu()
                choice = input("Enter your choice (1-7): ")

                if choice == '1':
                    self.automate_all()
                elif choice == '2':
                    self.extract_integration()
                elif choice == '3':
                    self.generate_screenshots()
                elif choice == '4':
                    self.generate_ai_descriptions()
                elif choice == '5':
                    self.generate_documentation()
                elif choice == '6':
                    self.clean_all_generated_files()
                elif choice == '7':
                    self.log("Exiting Celigo Automation. Goodbye!", color=Fore.CYAN)
                    break
                else:
                    self.log("Invalid choice. Please try again.", color=Fore.YELLOW)

        except Exception as e:
            self.log(f"An unexpected error occurred: {str(e)}", level="ERROR", color=Fore.RED)
        finally:
            if self.driver:
                self.log("Closing the browser...", color=Fore.CYAN)
                self.driver.quit()

    def automate_all(self):
        self.log("Starting full automation process...", color=Fore.CYAN)
        self.extract_integration()
        self.generate_screenshots()
        self.clean_integration_data()
        self.generate_ai_descriptions()
        self.process_scripts()
        self.generate_documentation()
        self.log("Full automation process completed!", color=Fore.GREEN)

    def extract_integration(self):
        self.log("Extracting integration...", color=Fore.CYAN)
        self.integration_id, self.integration_data, self.flow_ids = self.process_integration_zip()
        if self.integration_id and self.integration_data:
            self.log("Integration extracted successfully.", color=Fore.GREEN)
        else:
            self.log("Failed to extract integration.", color=Fore.RED)

    def generate_screenshots(self):
        self.log("Generating screenshots...", color=Fore.CYAN)
        if self.integration_id and hasattr(self, 'integration_data'):
            urls = self.generate_urls(self.integration_id, self.integration_data)
            for flow_id in self.flow_ids:
                self.capture_flow_screenshot(self.integration_id, flow_id)
            self.take_screenshots(urls)
            self.log("Screenshots generated successfully.", color=Fore.GREEN)
        else:
            self.log("Integration data not available. Please extract integration first.", color=Fore.YELLOW)

    def clean_all_generated_files(self):
        self.log("Cleaning all generated files...", color=Fore.CYAN)
        try:
            # Remove the extracted directory
            extracted_dir = os.path.join(self.celigo_ai_dir, "extracted")
            if os.path.exists(extracted_dir):
                shutil.rmtree(extracted_dir)
                self.log(f"Removed extracted directory: {extracted_dir}", color=Fore.GREEN)

            # Remove the output directory
            if os.path.exists(self.output_directory):
                shutil.rmtree(self.output_directory)
                self.log(f"Removed output directory: {self.output_directory}", color=Fore.GREEN)

            # Remove any generated .docx files in the CeligoAI directory
            for file in os.listdir(self.celigo_ai_dir):
                if file.endswith(".docx"):
                    os.remove(os.path.join(self.celigo_ai_dir, file))
                    self.log(f"Removed generated document: {file}", color=Fore.GREEN)

            self.log("All generated files cleaned successfully.", color=Fore.GREEN)
        except Exception as e:
            self.log(f"Error while cleaning generated files: {str(e)}", level="ERROR", color=Fore.RED)

            # Keep the browser open until user decides to close
            input("Press Enter to close the browser and end the automation...")
        except Exception as e:
            self.log(f"An unexpected error occurred: {str(e)}", level="ERROR", color=Fore.RED)
        finally:
            if self.driver:
                self.log("Closing the browser...", color=Fore.CYAN)
                self.driver.quit()

if __name__ == "__main__":
    automation = CeligoAutomation()
    automation.perform_automation()